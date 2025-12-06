# modules/text_extraction.py
from pathlib import Path
from typing import Optional
from .utils import safe_write_text, clean_extracted_text

# optional imports; errors will surface at runtime if missing
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx
except Exception:
    docx = None

def _extract_pdf_from_stream(file_stream) -> str:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber not installed in this venv. pip install pdfplumber")
    file_stream.seek(0)
    texts = []
    with pdfplumber.open(file_stream) as pdf:
        for p in pdf.pages:
            t = p.extract_text()
            texts.append(t or "")
    return "\n\n".join(texts)

def _extract_docx_from_stream(file_stream) -> str:
    if docx is None:
        raise RuntimeError("python-docx not installed. pip install python-docx")
    import tempfile
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.write(file_stream.read())
    tmp.flush()
    tmp.close()
    doc = docx.Document(tmp.name)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    Path(tmp.name).unlink(missing_ok=True)
    return "\n\n".join(paras)

def _extract_txt_from_stream(file_stream) -> str:
    file_stream.seek(0)
    raw = file_stream.read()
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="ignore")
    return raw

def extract_text(uploaded, save_to: Optional[str] = None) -> str:
    # uploaded is a file-like object from Streamlit or a path string
    name = uploaded.name.lower() if hasattr(uploaded, "name") else str(uploaded).lower()
    if name.endswith(".pdf"):
        raw = _extract_pdf_from_stream(uploaded)
    elif name.endswith(".docx") or name.endswith(".doc"):
        raw = _extract_docx_from_stream(uploaded)
    elif name.endswith(".txt"):
        raw = _extract_txt_from_stream(uploaded)
    else:
        raise ValueError("Unsupported file type")
    cleaned = clean_extracted_text(raw)
    if save_to:
        safe_write_text(save_to, cleaned)
    return cleaned
# inside text_extraction.py - add docx extraction function
import docx
def _extract_docx_from_stream(file_stream):
    import tempfile
    tf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tf.write(file_stream.read())
    tf.flush(); tf.close()
    doc = docx.Document(tf.name)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    Path(tf.name).unlink(missing_ok=True)
    return "\n\n".join(paras)
